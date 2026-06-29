from pathlib import Path
from docx import Document

p = Path(r"c:\Users\pavan\OneDrive\Documents\Pavani Kengana_Testing Profile.docx")
doc = Document(str(p))
align_map = {0: "LEFT", 1: "CENTER", 2: "RIGHT", 3: "JUSTIFY", None: "None"}

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip().replace("\n", " ")
    if not text and not para.runs:
        continue
    al = align_map.get(para.alignment, para.alignment)
    style = para.style.name if para.style else ""
    print(f"{i:3} | {al:8} | {style:20} | {text[:100]!r}")

print("\n=== TABLES ===")
print("count", len(doc.tables))
for ti, table in enumerate(doc.tables):
    print(f"Table {ti}: {len(table.rows)}x{len(table.columns)}")
    for ri, row in enumerate(table.rows):
        cells = [c.text.strip()[:35].replace("\n", " ") for c in row.cells]
        print(f"  r{ri}: {cells}")
