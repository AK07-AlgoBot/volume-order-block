from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

p = Path(r"c:\Users\pavan\OneDrive\Documents\Pavani Kengana_Testing Profile.docx")
doc = Document(str(p))
table = doc.tables[0]

# unique cells in row 3
row = table.rows[3]
seen = set()
cells = []
for cell in row.cells:
    if id(cell._tc) not in seen:
        seen.add(id(cell._tc))
        cells.append(cell)

print("Unique cells in row 3:", len(cells))
for i, cell in enumerate(cells):
    print(f"\nCell {i}: {len(cell.paragraphs)} paragraphs, {len(cell.text)} chars")
    for j, para in enumerate(cell.paragraphs[:15]):
        t = para.text.strip()
        if t:
            print(f"  p{j} align={para.alignment}: {t[:70]!r}")

# Save full text for review
out = Path(r"c:\Users\pavan\arun\volume-order-block\scripts\_docx_extract.txt")
parts = []
for ri, row in enumerate(table.rows):
    seen = set()
    for cell in row.cells:
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        if cell.text.strip():
            parts.append(f"=== ROW {ri} ===\n{cell.text}\n")
out.write_text("\n".join(parts), encoding="utf-8")
print("\nWrote", out)
