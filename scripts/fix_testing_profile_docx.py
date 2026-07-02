"""Rebuild Pavani Kengana Testing Profile resume — section-wise layout (no page overlap)."""

from __future__ import annotations

import re
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(r"c:\Users\pavan\OneDrive\Documents")
SRC = DOCS / "Pavani Kengana_Testing Profile.docx"
BACKUP = DOCS / "Pavani Kengana_Testing Profile_backup.docx"
OUT = DOCS / "Pavani Kengana_Testing Profile.docx"
OUT_NEW = DOCS / "Pavani Kengana_Testing Profile_fixed.docx"

NAVY = RGBColor(0x1F, 0x49, 0x7D)
GRAY = RGBColor(0x55, 0x55, 0x55)

SECTION_HEADERS = {
    "PROFILE SUMMARY",
    "CORE COMPETENCIES",
    "CERTIFICATIONS",
    "TECHNICAL SKILLS",
    "SOFT SKILLS",
    "PROFESSIONAL EXPERIENCE",
    "WORK EXPERIENCE",
    "CAREER TIMELINE",
    "EDUCATION",
    "PERSONAL DETAILS",
    "KEY ACHIEVEMENTS",
    "PROJECT HIGHLIGHTS",
}

SOFT_SKILLS = [
    "Strong communication and stakeholder coordination",
    "Independent ownership with collaborative team delivery",
    "Adaptable in agile, fast-changing project environments",
    "Mentoring and onboarding support for new team members",
]

PROFILE_SUMMARY_SHORT = (
    "Performance-driven QA professional with 10+ years across manual and automation testing. "
    "Experienced in STLC ownership from requirements through closure, with strong exposure to "
    "retail, telecom, and insurance domains at TCS, Foray, Capgemini, and Infosys."
)

OBJECTIVE_POLISHED = (
    "Performance-driven Software Testing professional with a strong track record of quality delivery, "
    "seeking challenging QA leadership roles with a reputed organization, preferably in Bangalore."
)

CORE_COMPETENCIES_SHORT = [
    "10+ years in Manual & Automation testing (Selenium WebDriver) across full STLC",
    "Requirement analysis, test planning, TestNG scripting, execution, and closure",
    "Defect management with JIRA, qTest, and Zephyr",
    "Data-driven & POM frameworks; Jenkins, Maven, GitHub CI/CD pipelines",
    "Functional, regression, integration, API (Postman), UI, and accessibility testing",
    "Independent delivery in agile teams across retail, telecom, and insurance domains",
]


def unique_row_cells(row):
    seen, out = set(), []
    for cell in row.cells:
        tid = id(cell._tc)
        if tid not in seen:
            seen.add(tid)
            out.append(cell)
    return out


def cell_text(cell) -> str:
    lines = [p.text.rstrip() for p in cell.paragraphs]
    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if not line.strip():
            blank_run += 1
            if blank_run <= 1:
                cleaned.append("")
            continue
        blank_run = 0
        cleaned.append(line.strip())
    return "\n".join(cleaned).strip()


def is_section_header(text: str) -> bool:
    t = re.sub(r"\s+", " ", text.strip().rstrip(":"))
    if t.upper() in SECTION_HEADERS:
        return True
    return len(t) < 45 and t == t.upper() and not t.endswith(".")


def normalize_header(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().rstrip(":")).upper()


def split_sections(text: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current = "_preamble"
    sections[current] = []
    for raw in text.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if is_section_header(line):
            current = normalize_header(line)
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return sections


def pick_source() -> Path:
    if BACKUP.exists():
        try:
            doc = Document(str(BACKUP))
            if doc.tables:
                row3 = unique_row_cells(doc.tables[0].rows[3])
                if len(row3) >= 5 and len(cell_text(row3[4])) > 500:
                    return BACKUP
        except Exception:
            pass
    return SRC


def extract_content(source: Path) -> tuple[str, list[str], str, str, str]:
    doc = Document(str(source))
    table = doc.tables[0]
    row0 = unique_row_cells(table.rows[0])
    header_text = cell_text(row0[-1]) if row0 else "PAVANI KENGANA"

    contact_parts: list[str] = []
    for cell in unique_row_cells(table.rows[1]):
        t = cell_text(cell)
        if t:
            contact_parts.append(t.replace("\n", " ").strip())

    objective = ""
    for cell in unique_row_cells(table.rows[2]):
        t = cell_text(cell)
        if t and not is_section_header(t):
            objective = t
            break

    row3 = unique_row_cells(table.rows[3])
    left_text = cell_text(row3[1]) if len(row3) >= 5 else ""
    right_text = cell_text(row3[4]) if len(row3) >= 5 else ""

    return header_text, contact_parts, objective, left_text, right_text


def parse_left_sidebar(left_text: str) -> dict[str, list[str]]:
    sections = split_sections(left_text)
    return {
        "certifications": sections.get("CERTIFICATIONS", []),
        "technical_skills": sections.get("TECHNICAL SKILLS", []),
        "education": sections.get("EDUCATION", []),
        "personal_details": sections.get("PERSONAL DETAILS", []),
    }


def is_role_line(line: str) -> bool:
    return bool(re.match(r"^(Since\s+)?[A-Za-z]{3}['\u2019]?\d{2}", line))


def is_project_line(line: str) -> bool:
    upper = line.upper()
    return upper.startswith("PROJECTS UNDERTAKEN") or (
        "| CLIENT:" in upper or upper.startswith("CLIENT:")
    )


def format_work_experience(right_text: str) -> list[str]:
    sections = split_sections(right_text)
    raw_lines = sections.get("WORK EXPERIENCE", [])
    if not raw_lines:
        for key, lines in sections.items():
            if key not in {"PROFILE SUMMARY", "CAREER TIMELINE", "_preamble"}:
                raw_lines.extend(lines)

    blocks: list[str] = []
    current_role: list[str] = []
    current_bullets: list[str] = []
    in_kra = False
    seen_bullets: set[str] = set()

    def flush_role():
        nonlocal current_role, current_bullets, in_kra, seen_bullets
        if not current_role and not current_bullets:
            return
        blocks.extend(current_role)
        if current_bullets:
            blocks.append("Key Result Areas:")
            blocks.extend(f"• {b}" for b in current_bullets)
            blocks.append("")
        current_role = []
        current_bullets = []
        in_kra = False
        seen_bullets = set()

    for line in raw_lines:
        stripped = line.strip()
        if not stripped:
            continue
        upper = normalize_header(stripped)
        if upper == "KEY RESULT AREAS":
            in_kra = True
            continue
        if is_role_line(stripped) or is_project_line(stripped):
            flush_role()
            current_role.append(stripped)
            continue
        if in_kra:
            key = re.sub(r"\s+", " ", stripped.lower())
            if key not in seen_bullets:
                seen_bullets.add(key)
                current_bullets.append(stripped)
        else:
            current_role.append(stripped)

    flush_role()
    return blocks


def build_sidebar_text(sidebar: dict[str, list[str]]) -> str:
    parts: list[str] = ["CERTIFICATIONS", ""]
    parts.extend(sidebar["certifications"] or ["ISTQB Certification"])
    parts.extend(["", "TECHNICAL SKILLS", ""])
    parts.extend(sidebar["technical_skills"])
    parts.extend(["", "EDUCATION", ""])
    parts.extend(sidebar["education"])
    parts.extend(["", "PERSONAL DETAILS", ""])
    parts.extend(sidebar["personal_details"])
    return "\n".join(parts).strip()


def build_page1_right_column() -> str:
    parts = [
        "PROFILE SUMMARY",
        "",
        PROFILE_SUMMARY_SHORT,
        "",
        "CORE COMPETENCIES",
        "",
    ]
    parts.extend(f"• {item}" for item in CORE_COMPETENCIES_SHORT)
    parts.extend(["", "SOFT SKILLS", ""])
    parts.extend(f"• {item}" for item in SOFT_SKILLS)
    return "\n".join(parts).strip()


def set_cell_margins(cell, *, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for tag, val in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_run(paragraph, text: str, *, size=10.5, bold=False, italic=False, color=GRAY):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def set_paragraph_bottom_border(paragraph, *, color: str = "4472C4", size: int = 16):
    """Full-width horizontal rule (size = eighths of a point; 16 ≈ 2pt bold line)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_full_width_rule(cell, *, space_before: int = 0, space_after: int = 0, first: bool = False):
    p = cell.paragraphs[0] if first and cell.paragraphs else cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    set_paragraph_bottom_border(p)
    return p


def fill_objective_block(cell, text: str):
    cell.text = ""

    add_full_width_rule(cell, space_before=2, space_after=6, first=True)

    mid = cell.add_paragraph()
    mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mid.paragraph_format.space_before = Pt(0)
    mid.paragraph_format.space_after = Pt(0)
    add_run(mid, text, size=10, italic=True)

    add_full_width_rule(cell, space_before=6, space_after=2)


def add_section_header(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title.upper(), size=11, bold=True, color=NAVY)
    return p


def add_body_paragraph(doc: Document, text: str, *, size=10.5, bold=False, bullet=False, role=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    if bullet:
        p.paragraph_format.left_indent = Inches(0.18)
        text = text if text.startswith("•") else f"• {text}"
    if role:
        p.paragraph_format.space_before = Pt(6)
        add_run(p, text, size=10.5, bold=True, color=NAVY)
        return p
    add_run(p, text, size=size, bold=bold)
    return p


def fill_cell_lines(cell, text: str, *, body_size=10.5):
    cell.text = ""
    first = True
    for raw_line in text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            if first:
                continue
            cell.add_paragraph()
            continue

        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(2)

        if is_section_header(line):
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            add_run(p, line.strip().upper(), size=11, bold=True, color=NAVY)
            continue

        if line.lstrip().startswith("•"):
            p.paragraph_format.left_indent = Inches(0.12)
            add_run(p, line.strip(), size=body_size)
            continue

        add_run(p, line.strip(), size=body_size)


def add_work_experience_section(doc: Document, work_lines: list[str]):
    add_section_header(doc, "Work Experience")
    for line in work_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == "Key Result Areas:":
            add_body_paragraph(doc, stripped, bold=True)
            continue
        if stripped.lstrip().startswith("•"):
            add_body_paragraph(doc, stripped.lstrip("• ").strip(), bullet=True)
            continue
        if is_role_line(stripped):
            add_body_paragraph(doc, stripped, role=True)
            continue
        add_body_paragraph(doc, stripped)


def main() -> None:
    source = pick_source()
    header_text, contact_parts, objective, left_text, right_text = extract_content(source)

    if not BACKUP.exists() and SRC.exists():
        shutil.copy2(SRC, BACKUP)

    sidebar = parse_left_sidebar(left_text)
    sidebar_out = build_sidebar_text(sidebar)
    page1_right = build_page1_right_column()
    work_lines = format_work_experience(right_text)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    # Page-1 intro only: header + contact + objective + short two-column block.
    intro = doc.add_table(rows=4, cols=2)
    intro.autofit = False
    intro.columns[0].width = Inches(2.05)
    intro.columns[1].width = Inches(4.65)

    header = intro.rows[0].cells[0].merge(intro.rows[0].cells[1])
    lines = [ln.strip() for ln in header_text.split("\n") if ln.strip()]
    name = lines[0] if lines else "PAVANI KENGANA"
    title = lines[1] if len(lines) > 1 else "SOFTWARE TESTING & QA PROFESSIONAL"
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(4)
    r1 = hp.add_run(name.upper())
    r1.bold = True
    r1.font.size = Pt(20)
    r1.font.color.rgb = NAVY
    hp.add_run("\n")
    r2 = hp.add_run(title.upper())
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = NAVY

    contact = intro.rows[1].cells[0].merge(intro.rows[1].cells[1])
    cp = contact.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(6)
    cr = cp.add_run("  |  ".join(contact_parts))
    cr.font.size = Pt(9.5)
    cr.font.color.rgb = GRAY

    obj = intro.rows[2].cells[0].merge(intro.rows[2].cells[1])
    fill_objective_block(obj, OBJECTIVE_POLISHED)

    left = intro.rows[3].cells[0]
    right = intro.rows[3].cells[1]
    shade_cell(left, "F2F6FB")
    set_cell_margins(left)
    set_cell_margins(right)
    fill_cell_lines(left, sidebar_out, body_size=10)
    fill_cell_lines(right, page1_right, body_size=10)

    # Work experience only below the table — full width, no sidebar bleed on page 2+.
    add_work_experience_section(doc, work_lines)

    for target in (OUT_NEW, OUT):
        try:
            doc.save(str(target))
            saved = target
            break
        except PermissionError:
            continue
    else:
        saved = OUT_NEW
        doc.save(str(saved))

    print(f"Rebuilt resume: {saved}")
    print(f"Source content: {source.name}")
    print(f"Sidebar: {len(sidebar_out)} chars | Page-1 right: {len(page1_right)} chars")
    print(f"Work experience lines: {len(work_lines)}")
    print(f"Layout: page-1 two-column (summary + competencies) + full-width work history")
    print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    main()
